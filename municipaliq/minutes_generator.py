"""Generates MGL-compliant draft meeting minutes as a .docx file."""
import json
import os
import re
import sys
from pathlib import Path
from typing import Optional

from openai import OpenAI
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OLLAMA_HOST = os.environ.get('OLLAMA_HOST', 'http://localhost:11434')
_MODEL = os.environ.get('OLLAMA_MODEL', 'qwen3:1.7b')

_ROOT = Path(__file__).parent.parent
_NAME_CORRECTIONS_PATH = _ROOT / 'data' / 'name_corrections.json'

# Stage 1: chunk size for extraction pass (~3K tokens, well within context)
_CHUNK_CHARS = 12_000
_BODY_FONT_PT = 12
_FONT_NAME = 'Times New Roman'
_SECTION_SPACE_PT = 12  # space before ALL-CAPS section headings
_PARA_SPACE_PT = 6      # space_after for body paragraphs
_PROSE_MARKER = '[prose]'
_UNCLEAR = '[UNCLEAR]'
_MEMBERS_PRESENT_FALLBACK = f'{_UNCLEAR} — list Select Board members present'
_DEFAULT_LOCATION = 'Hardwick Town Hall'
_MSG_ROLE = 'role'
_MSG_CONTENT = 'content'

# Strips <think>...</think> blocks that reasoning models emit.
_THINK_RE = re.compile(r'<think>.*?</think>', re.DOTALL)

_EXTRACT_SYSTEM = """\
You are reading a portion of a town meeting transcript. Extract only the \
factual content relevant to official meeting minutes. Do not invent anything.

From this excerpt, list:
- Topics or agenda items discussed (brief phrase)
- Key points made during each discussion
- Any votes taken (motion text, who moved, who seconded, outcome)
- Any decisions or actions agreed upon
- Which Select Board members were present (use the known names list below)
- Start or end times if mentioned

Transcript format notes:
- ">>" marks a speaker change; speakers are not always named.
- Votes are often informal voice votes: "All in favor?" followed by "I" responses \
means the motion passed unanimously. Always record these as votes.
- Motion text may span several transcript lines — piece it together.

{board_context}\
Do NOT list presenters, guests, or members of the public as attendees — \
only Select Board members (and Finance Committee members if this is a joint meeting).
Be concise and factual. If nothing relevant is in this excerpt, say so briefly.
"""

# Stage 2 system: model fills [prose] markers in a pre-built skeleton.
_FILL_SYSTEM = """\
You are a professional municipal clerk completing a meeting minutes document.

Your ONLY output must be the skeleton below, with every [prose] marker replaced \
by factual prose sentences. Do not output the skeleton and then a separate \
"completed" version — output the final completed document once, in full.

Rules:
- Replace each [prose] with one or more formal prose sentences in third-person past tense.
- Do NOT change any text outside [prose] markers (headings, pre-filled lines stay as-is).
- No bullet points, no markdown, no asterisks.
- If you have no facts for a section, write "No discussion was recorded for this item."
- For votes, end the section with a standalone sentence: [Name] moved to [motion text]. \
The motion was seconded by [Name] and [passed/failed] [unanimously/X–Y] via voice vote. \
If the vote was a formal roll-call, write "via roll-call vote" instead.
- Write [UNCLEAR] only for genuinely unknown specifics.
- Output nothing before the first line of the skeleton and nothing after the last line.
"""

_FILL_USER = """\
MEETING DATE: {date}
MEETING TIME: {time}
LOCATION: {location}

{board_context}\
EXTRACTED FACTS:
{facts}

Skeleton to complete:

{skeleton}"""


def _load_name_corrections() -> dict:
    if not _NAME_CORRECTIONS_PATH.exists():
        return {}
    return json.loads(_NAME_CORRECTIONS_PATH.read_text())


def _apply_name_corrections(text: str, corrections: dict) -> str:
    for wrong, right in corrections.items():
        text = text.replace(wrong, right)
    return text


def _strip_thinking(text: str) -> str:
    text = _THINK_RE.sub('', text)
    # Prefill can cut off the opening <think> tag, leaving a stray </think>.
    # Everything before it is internal reasoning — discard it.
    if '</think>' in text:
        text = text.split('</think>', 1)[-1]
    return text.strip()


def _chunk(text: str, size: int) -> list[str]:
    positions = range(0, len(text), size)
    return [text[pos:pos + size] for pos in positions]


def _matches_name(member: str, role_name: str) -> bool:
    """Return True if member and role_name refer to the same person.

    Uses prefix matching so every word in role_name must match a word in the
    member name.  Handles missing middle initials ('Eric Vollheim' matches
    'Eric W. Vollheim') and nicknames ('Jeff' matches 'Jeffrey').
    """
    if not role_name:
        return False
    member_words = member.lower().split()
    for rw in role_name.lower().split():
        found = False
        for mw in member_words:
            if mw.startswith(rw) or rw.startswith(mw):
                found = True
        if not found:
            return False
    return True


def _ordered_members(members: list[str], chair: str, clerk: str) -> list[str]:
    """Return members sorted Chair first, Clerk last, others in between."""
    chairs = [name for name in members if _matches_name(name, chair)]
    clerks = [name for name in members if _matches_name(name, clerk) and name not in chairs]
    middle = [name for name in members if name not in chairs and name not in clerks]
    return chairs + middle + clerks


def _filter_absent(members: list[str], absent: list[str]) -> list[str]:
    """Remove absent members from the present list using fuzzy name matching."""
    if not absent:
        return list(members)
    return [
        member
        for member in members
        if not any(_matches_name(member, absentee) for absentee in absent)
    ]


def _board_context(board_info: dict) -> str:
    """Build a short prompt snippet listing known board members by name."""
    members = board_info.get('members') or []
    chair = board_info.get('chair') or ''
    if not members:
        return ''
    chair_note = f' ({chair} is the Chair)' if chair else ''
    ordered = _ordered_members(members, chair, board_info.get('clerk') or '')
    members_str = ', '.join(ordered)
    return (
        f'Known Select Board members{chair_note}: {members_str}. '
        'Use these exact names when identifying board members.\n'
    )


def _extract_chunk(
    client: OpenAI, system: str, chunk: str, chunk_num: int, total: int,
) -> str:
    label = f'Transcript excerpt {chunk_num} of {total}:\n\n{chunk}'
    response = client.chat.completions.create(
        model=_MODEL,
        messages=[
            {_MSG_ROLE: 'system', _MSG_CONTENT: system},
            {_MSG_ROLE: 'user', _MSG_CONTENT: label},
        ],
    )
    return _strip_thinking(response.choices[0].message.content)


def _extract_facts(client: OpenAI, transcript: str, board_info: dict) -> str:
    """Stage 1: extract key facts from each transcript chunk."""
    system = _EXTRACT_SYSTEM.format(board_context=_board_context(board_info))
    chunks = _chunk(transcript, _CHUNK_CHARS)
    total = len(chunks)
    parts = []
    for num, chunk in enumerate(chunks, 1):
        excerpt = _extract_chunk(client, system, chunk, num, total)
        parts.append(f'[Excerpt {num}]\n{excerpt}')
    return '\n\n'.join(parts)


# Heading-like lines: short (≤80 chars), not starting with tab/spaces or digits
_AGENDA_HEADING_RE = re.compile(r'^(?!\s)(?!\d)(?!https?://).{3,80}$')
# Marks executive session block
_EXEC_RE = re.compile(r'executive session', re.IGNORECASE)
# MGL citation pattern
_MGL_RE = re.compile(
    r'(MGL?\s+(?:c\.|ch\.?|chapter)?\s*30A.*?(?:\n|$)(?:.*?(?:\n|$))?)',
    re.IGNORECASE,
)


_SUB_ITEM_RE = re.compile(r'^[a-zA-Z0-9][.)]\s')
_METADATA_RE = re.compile(
    r'^(http|Streaming|Posted|Scheduled|Last Modified)', re.IGNORECASE,
)
_BARE_TIME_RE = re.compile(r'^\d{1,2}:\d{2}')


def _is_agenda_section(stripped: str) -> bool:
    if _SUB_ITEM_RE.match(stripped):
        return False
    if _METADATA_RE.match(stripped):
        return False
    if _BARE_TIME_RE.match(stripped):
        return False
    return bool(_AGENDA_HEADING_RE.match(stripped))


def _agenda_sections(agenda_text: str) -> list[str]:
    """Return top-level section titles from agenda text.

    Skips the header block (before the first blank line) and filters
    sub-items and metadata lines.
    """
    sections = []
    passed_header = False
    for line in agenda_text.split('\n'):
        stripped = line.strip()
        if not stripped:
            passed_header = True
            continue
        if not passed_header:
            continue
        if _is_agenda_section(stripped):
            sections.append(stripped)
    return sections


def _append_exec_line(exec_lines: list[str], stripped: str) -> bool:
    """Append stripped line; return True if the exec block has ended."""
    if not stripped:
        return bool(exec_lines)
    exec_lines.append(stripped)
    return False


def _collect_exec_lines(agenda_text: str) -> list[str]:
    exec_lines: list[str] = []
    capturing = False
    for line in agenda_text.split('\n'):
        stripped = line.strip()
        capturing = capturing or bool(_EXEC_RE.search(stripped))
        if capturing and _append_exec_line(exec_lines, stripped):
            break
    return exec_lines


def _exec_session_prose(agenda_text: str) -> str:
    """Generate executive session minutes text from agenda language."""
    exec_lines = _collect_exec_lines(agenda_text)
    if not exec_lines:
        return _PROSE_MARKER
    return _format_exec_prose(' '.join(exec_lines))


_MGL_CITE_RE = re.compile(
    r'(MGL?\s+(?:c\.|ch\.?|chapter)?\s*30A[^.]+\.)', re.IGNORECASE,
)
_DEFAULT_MGL = 'MGL Chapter 30A, Section 21'
_EXEC_TAIL = (
    ' The Board did not return to open session following the Executive Session.'
    ' No action was taken in Executive Session that required a vote in open session.'
)


def _format_exec_prose(purpose_text: str) -> str:
    mgl_match = _MGL_CITE_RE.search(purpose_text)
    mgl_cite = mgl_match.group(1).strip() if mgl_match else _DEFAULT_MGL
    after_cite = purpose_text[mgl_match.end():].strip(' -–') if mgl_match else purpose_text
    intro = (
        f'Prior to the Open Meeting, the Select Board voted to enter Executive Session '
        f'pursuant to {mgl_cite}.'
    )
    purpose = f' The stated purpose was: {after_cite}' if after_cite else ''
    return intro + purpose + _EXEC_TAIL


def _build_skeleton(meeting: dict, agenda_text: str, board_info: dict) -> str:
    """Build a pre-structured minutes skeleton with [prose] placeholders.

    The header is filled with actual meeting values. Executive session is
    generated from agenda language. All other sections get [prose] markers
    for the model to fill in.
    """
    members = board_info.get('members') or []
    chair = board_info.get('chair') or ''
    clerk = board_info.get('clerk') or ''
    ordered = _ordered_members(members, chair, clerk) if members else []
    absent = meeting.get('members_absent') or []
    members_line = ', '.join(_filter_absent(ordered, absent)) or _MEMBERS_PRESENT_FALLBACK
    lines = [f'MEMBERS PRESENT: {members_line}', '']
    if absent:
        absent_display = ', '.join(absent)
        lines += [f'MEMBERS ABSENT: {absent_display}', '']
    lines += _skeleton_body(agenda_text)
    lines += ['ADJOURNMENT', _PROSE_MARKER]
    return '\n'.join(lines)


_TIME_ANNOTATION_RE = re.compile(r'\s*[-–]\s*\d{1,2}:\d{2}\s*(AM|PM).*$')


def _skeleton_body(agenda_text: str) -> list[str]:
    if not agenda_text:
        return ['CALL TO ORDER', _PROSE_MARKER, '', 'BUSINESS', _PROSE_MARKER, '']
    lines: list[str] = []
    for section in _agenda_sections(agenda_text):
        title_upper = _TIME_ANNOTATION_RE.sub('', section.upper()).strip()
        if 'OPEN SESSION' in title_upper or 'ADJOURNMENT' in title_upper:
            continue
        if _EXEC_RE.search(section):
            lines.extend([title_upper, _exec_session_prose(agenda_text), ''])
        else:
            lines.extend([title_upper, _PROSE_MARKER, ''])
    return lines


_DIVIDER_RE = re.compile(
    r'\n(?:Finalized|Completed|Final|Here is the completed)[^\n]*\n',
    re.IGNORECASE,
)
_SKELETON_HEADER_RE = re.compile(
    r'^HARDWICK SELECT BOARD\s*\nMeeting Minutes\s*\n[^\n]+\n',
    re.MULTILINE,
)
# Strips "MEETING MINUTES" when the model repeats it as a section heading inside
# the body — it is already shown in the document header so it looks redundant there.
_MEETING_MINUTES_RE = re.compile(r'^MEETING MINUTES\s*$', re.MULTILINE | re.IGNORECASE)


def _clean_minutes_text(text: str, chair: str) -> str:
    text = _strip_thinking(text)
    text = text.replace('**', '')
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    divider_match = _DIVIDER_RE.search(text)
    if divider_match:
        text = text[divider_match.end():]
    text = _SKELETON_HEADER_RE.sub('', text)
    text = _MEETING_MINUTES_RE.sub('', text)
    text = re.sub(r'^\[prose\]\s*$', '', text, flags=re.MULTILINE)
    text = re.sub(r'\[specific [^\]]+\]', _UNCLEAR, text)
    text = re.sub(r"\[Chairperson's Name\]", chair or _UNCLEAR, text)
    text = re.sub(r'\[Your Name[^\]]*\]', '', text)
    return re.sub(r'\[date\]', '', text, flags=re.IGNORECASE)


def _write_minutes(
    client: OpenAI,
    meeting: dict,
    facts: str,
    board_info: dict,
    agenda_text: str,
) -> str:
    """Stage 2: fill a pre-built skeleton with prose from extracted facts."""
    skeleton = _build_skeleton(meeting, agenda_text, board_info)
    response = client.chat.completions.create(
        model=_MODEL,
        messages=[
            {_MSG_ROLE: 'system', _MSG_CONTENT: _FILL_SYSTEM},
            {_MSG_ROLE: 'user', _MSG_CONTENT: _FILL_USER.format(
                date=meeting.get('date', 'Unknown'),
                time=meeting.get('time', _UNCLEAR),
                location=meeting.get('location') or _DEFAULT_LOCATION,
                board_context=_board_context(board_info),
                facts=facts,
                skeleton=skeleton,
            )},
        ],
    )
    return _clean_minutes_text(
        response.choices[0].message.content,
        board_info.get('chair', ''),
    )


_SEPARATOR_RE = re.compile(r'^-{3,}$')
_BLANK_RUNS_RE = re.compile(r'\n{3,}')
# Numbered agenda item: "1. SOME TITLE" — kept as fallback in case model numbers headings
_NUMBERED_HEADING_RE = re.compile(r'^\d+\.\s+[A-Z]')
# Standalone ALL-CAPS line (heading with no number): at least 3 chars, no lowercase
_ALLCAPS_LINE_RE = re.compile(r'^[A-Z][A-Z0-9 /\-\(\)]{2,}$')
# Vote sentence: contains "seconded by" (Belmont-style motion format)
_VOTE_PARA_RE = re.compile(r'\bseconded\s+by\b', re.IGNORECASE)
_DRAFT_FONT_PT = 10
_HEADING_FONT_PT = 14


def _add_plain_paragraph(doc: Document, text: str, italic: bool = False) -> None:
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(_PARA_SPACE_PT)
    for run in para.runs:
        run.font.name = _FONT_NAME
        run.font.size = Pt(_BODY_FONT_PT)
        if italic:
            run.italic = True


def _add_docx_header(doc: Document, date: str, location: str) -> None:
    header_rows = (
        ('Minutes of the', False, _BODY_FONT_PT),
        ('HARDWICK SELECT BOARD MEETING', True, _HEADING_FONT_PT),
        (location or _DEFAULT_LOCATION, False, _BODY_FONT_PT),
        (date, False, _BODY_FONT_PT),
    )
    for text, bold, size in header_rows:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.bold = bold
        run.font.name = _FONT_NAME
        run.font.size = Pt(size)

    draft_para = doc.add_paragraph()
    draft_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    draft_run = draft_para.add_run('DRAFT — Subject to approval at the next regular meeting')
    draft_run.italic = True
    draft_run.font.name = _FONT_NAME
    draft_run.font.size = Pt(_DRAFT_FONT_PT)


def _add_section_heading(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(_SECTION_SPACE_PT)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.bold = True
    run.font.name = _FONT_NAME
    run.font.size = Pt(_BODY_FONT_PT)


def _add_docx_line(doc: Document, stripped: str) -> None:
    if not stripped or _SEPARATOR_RE.match(stripped):
        return
    if _NUMBERED_HEADING_RE.match(stripped) or _ALLCAPS_LINE_RE.match(stripped):
        _add_section_heading(doc, stripped)
    elif stripped.startswith('- ') or stripped.startswith('* '):
        _add_plain_paragraph(doc, stripped[2:])
    else:
        _add_plain_paragraph(doc, stripped, italic=bool(_VOTE_PARA_RE.search(stripped)))


def _build_docx(date: str, location: str, body_text: str) -> Document:
    doc = Document()
    _add_docx_header(doc, date, location)
    for line in body_text.split('\n'):
        _add_docx_line(doc, line.strip())
    doc.add_paragraph('')
    _add_plain_paragraph(doc, 'Respectfully submitted by,', italic=True)
    doc.add_paragraph('')
    _add_plain_paragraph(doc, '[Recording Secretary]', italic=True)
    _add_plain_paragraph(doc, 'Recording Secretary', italic=True)
    return doc


def generate_minutes(
    meeting: dict,
    agenda_text: str,
    transcript_text: str,
    output_path: Path,
    board_info: Optional[dict] = None,
) -> Path:
    """
    Generate draft minutes using a two-stage approach:
      1. Extract key facts from transcript chunks (with name corrections applied)
      2. Fill a pre-built skeleton with prose from those facts
    Saves the result as a .docx file at output_path.
    """
    client = OpenAI(base_url=f'{OLLAMA_HOST}/v1', api_key='ollama')
    board_info = board_info or {}
    corrections = _load_name_corrections()

    sys.stdout.write('    Extracting facts from transcript...\n')
    facts = _extract_facts(client, transcript_text, board_info)
    facts = _apply_name_corrections(facts, corrections)

    sys.stdout.write('    Writing minutes from extracted facts...\n')
    minutes_text = _write_minutes(client, meeting, facts, board_info, agenda_text)
    minutes_text = _apply_name_corrections(minutes_text, corrections)

    doc = _build_docx(
        meeting.get('date', ''),
        meeting.get('location') or 'Hardwick Town Hall',
        minutes_text,
    )
    doc.save(str(output_path))
    return output_path
