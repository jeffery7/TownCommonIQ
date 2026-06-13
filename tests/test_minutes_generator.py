from unittest.mock import MagicMock

import pytest

from towncommoniq import minutes_generator


MEETING = {
    'date': '2024-03-15',
    'time': '6:30 PM',
    'location': 'Hardwick Town Hall',
    'status': 'held',
}

BOARD_INFO = {
    'chair': 'Eric Vollheim',
    'members': ['Jeffrey S. Schaaf', 'Eric W. Vollheim', 'William F. Tinker'],
}

AGENDA_TEXT = 'TEST BOARD\nAgenda\n\nCall to Order\nApproval of Minutes\nAdjournment'
TRANSCRIPT_TEXT = 'Chair Vollheim called the meeting to order at 6:30 PM. All members present.'

MOCK_FACTS_TEXT = (
    'MEMBERS PRESENT: Jeffrey S. Schaaf, Eric W. Vollheim, William F. Tinker\n'
    'Meeting called to order at 6:30 PM.\n'
    'VOTED to adjourn. Vollheim moved, Schaaf seconded. All yes. 3-0.'
)

MOCK_MINUTES_TEXT = (
    'HARDWICK SELECT BOARD\n'
    'Meeting Minutes\n'
    '2024-03-15 | 6:30 PM | Hardwick Town Hall\n\n'
    'MEMBERS PRESENT: Jeffrey S. Schaaf, Eric W. Vollheim, William F. Tinker\n\n'
    '1. CALL TO ORDER\n'
    'Chair Vollheim called the meeting to order at 6:30 PM.\n\n'
    'ADJOURNMENT\n'
    'The meeting was adjourned at 7:00 PM.\n\n'
    'DRAFT — Subject to approval at the next regular meeting.\n'
)


def _make_openai_response(content: str) -> MagicMock:
    """Build a mock that mimics openai ChatCompletion response structure."""
    choice = MagicMock()
    choice.message.content = content
    response = MagicMock()
    response.choices = [choice]
    return response


@pytest.fixture()
def mock_openai(monkeypatch):
    mock_client = MagicMock()
    mock_client.chat.completions.create.side_effect = [
        _make_openai_response(MOCK_FACTS_TEXT),   # _extract_facts call
        _make_openai_response(MOCK_MINUTES_TEXT),  # _write_minutes call
    ]

    monkeypatch.setattr(
        'towncommoniq.minutes_generator.OpenAI',
        lambda **kw: mock_client,
    )
    return mock_client


class TestLoadNameCorrections:
    def test_returns_empty_when_no_file(self, monkeypatch, tmp_path):
        monkeypatch.setattr(minutes_generator, '_NAME_CORRECTIONS_PATH', tmp_path / 'none.json')
        assert minutes_generator._load_name_corrections() == {}

    def test_loads_from_file(self, monkeypatch, tmp_path):
        f = tmp_path / 'corrections.json'
        f.write_text('{"Mark Corsick": "Mark Korzec"}')
        monkeypatch.setattr(minutes_generator, '_NAME_CORRECTIONS_PATH', f)
        result = minutes_generator._load_name_corrections()
        assert result == {'Mark Corsick': 'Mark Korzec'}


class TestAgendaSections:
    def test_skips_sub_items(self):
        agenda = 'BOARD\n\nMain Topic\na. Sub-item one\nb) Sub-item two'
        sections = minutes_generator._agenda_sections(agenda)
        assert 'Main Topic' in sections
        assert not any('Sub-item' in s for s in sections)

    def test_skips_streaming_metadata(self):
        agenda = 'BOARD\n\nMain Topic\nStreaming: YouTube\nhttp://example.com'
        sections = minutes_generator._agenda_sections(agenda)
        assert 'Main Topic' in sections
        assert not any('Streaming' in s or 'http' in s for s in sections)

    def test_skips_bare_time_lines(self):
        agenda = 'BOARD\n\nMain Topic\n6:30 PM'
        sections = minutes_generator._agenda_sections(agenda)
        assert 'Main Topic' in sections
        assert not any(':' in s and 'PM' in s for s in sections)

    def test_skips_header_before_first_blank(self):
        agenda = 'Board Name\nDate: 2024-01-01\n\nCall to Order'
        sections = minutes_generator._agenda_sections(agenda)
        assert sections == ['Call to Order']


class TestExecSessionProse:
    def test_returns_prose_marker_when_no_exec_section(self):
        result = minutes_generator._exec_session_prose('Call to Order\nApproval of Minutes\nAdjournment')
        assert result == '[prose]'

    def test_uses_default_citation_when_no_mgl(self):
        result = minutes_generator._exec_session_prose('Executive Session to discuss personnel')
        assert 'MGL Chapter 30A, Section 21' in result

    def test_includes_mgl_citation_when_present(self):
        agenda = 'Executive Session pursuant to MGL c.30A §21(a)(1) to discuss litigation.'
        result = minutes_generator._exec_session_prose(agenda)
        assert 'MGL' in result
        assert '30A' in result

    def test_stops_capturing_at_blank_line(self):
        agenda = 'Executive Session to discuss personnel\n\nCall to Order'
        result = minutes_generator._exec_session_prose(agenda)
        assert 'Call to Order' not in result
        assert 'personnel' in result


class TestFilterAbsent:
    def test_removes_matching_member(self):
        members = ['Alice Smith', 'Eric W. Vollheim', 'Bob Jones']
        result = minutes_generator._filter_absent(members, ['Eric Vollheim'])
        assert result == ['Alice Smith', 'Bob Jones']

    def test_empty_absent_returns_all(self):
        members = ['Alice Smith', 'Bob Jones']
        assert minutes_generator._filter_absent(members, []) == ['Alice Smith', 'Bob Jones']

    def test_no_match_returns_all(self):
        members = ['Alice Smith', 'Bob Jones']
        assert minutes_generator._filter_absent(members, ['Charlie Smith']) == ['Alice Smith', 'Bob Jones']

    def test_removes_multiple_absent(self):
        members = ['Alice Smith', 'Eric W. Vollheim', 'Bob Jones']
        result = minutes_generator._filter_absent(members, ['Eric Vollheim', 'Bob Jones'])
        assert result == ['Alice Smith']


class TestBuildSkeletonSections:
    def test_open_session_omitted(self):
        agenda = 'BOARD\n\nOpen Session\nCall to Order\nAdjournment'
        skeleton = minutes_generator._build_skeleton({'date': '2024-03-15', 'time': '6:30 PM', 'location': 'Town Hall'}, agenda, {})
        assert 'OPEN SESSION' not in skeleton

    def test_adjournment_in_footer_not_agenda(self):
        agenda = 'BOARD\n\nCall to Order\nAdjournment'
        skeleton = minutes_generator._build_skeleton({'date': '2024-03-15', 'time': '6:30 PM', 'location': 'Town Hall'}, agenda, {})
        assert skeleton.count('ADJOURNMENT') == 1

    def test_exec_session_generates_prose(self):
        agenda = 'BOARD\n\nCall to Order\nExecutive Session to discuss personnel\nAdjournment'
        skeleton = minutes_generator._build_skeleton({'date': '2024-03-15', 'time': '6:30 PM', 'location': 'Town Hall'}, agenda, {})
        assert 'EXECUTIVE SESSION' in skeleton
        assert 'MGL' in skeleton

    def test_absent_member_excluded_from_present_line(self):
        board = {'members': ['Alice Smith', 'Eric W. Vollheim'], 'chair': '', 'clerk': ''}
        meeting = {'members_absent': ['Eric Vollheim']}
        skeleton = minutes_generator._build_skeleton(meeting, '', board)
        present_line = skeleton.split('\n')[0]
        assert 'Alice Smith' in present_line
        assert 'Eric' not in present_line

    def test_absent_line_shown_when_absent_set(self):
        board = {'members': ['Alice Smith', 'Eric W. Vollheim'], 'chair': '', 'clerk': ''}
        meeting = {'members_absent': ['Eric W. Vollheim']}
        skeleton = minutes_generator._build_skeleton(meeting, '', board)
        assert 'MEMBERS ABSENT: Eric W. Vollheim' in skeleton

    def test_no_absent_line_when_all_present(self):
        board = {'members': ['Alice Smith', 'Bob Jones'], 'chair': '', 'clerk': ''}
        skeleton = minutes_generator._build_skeleton({}, '', board)
        assert 'MEMBERS ABSENT' not in skeleton


class TestStripThinking:
    def test_removes_think_block(self):
        text = '<think>internal reasoning here</think>actual output'
        assert minutes_generator._strip_thinking(text) == 'actual output'

    def test_removes_multiline_think_block(self):
        text = '<think>\nline 1\nline 2\n</think>\nresult'
        assert minutes_generator._strip_thinking(text) == 'result'

    def test_handles_orphaned_close_tag(self):
        text = 'cut off reasoning\n</think>\nactual output'
        assert minutes_generator._strip_thinking(text) == 'actual output'

    def test_no_think_block_unchanged(self):
        text = 'plain output with no thinking'
        assert minutes_generator._strip_thinking(text) == 'plain output with no thinking'


class TestCleanMinutesText:
    def test_strips_standalone_meeting_minutes_heading(self):
        text = 'CALL TO ORDER\nSome prose.\n\nMEETING MINUTES\n\nMore prose.'
        result = minutes_generator._clean_minutes_text(text, 'Chair')
        assert 'MEETING MINUTES' not in result
        assert 'CALL TO ORDER' in result

    def test_strips_meeting_minutes_case_insensitive(self):
        text = 'Meeting Minutes\nSome prose.'
        result = minutes_generator._clean_minutes_text(text, 'Chair')
        assert 'Meeting Minutes' not in result


class TestApplyNameCorrections:
    def test_replaces_wrong_name(self):
        corrections = {'Mark Corsick': 'Mark Korzec'}
        result = minutes_generator._apply_name_corrections('Mark Corsick attended.', corrections)
        assert result == 'Mark Korzec attended.'

    def test_empty_corrections(self):
        result = minutes_generator._apply_name_corrections('No changes needed.', {})
        assert result == 'No changes needed.'


class TestOrderedMembers:
    _members = ['Jeffrey S. Schaaf', 'Eric W. Vollheim', 'William F. Tinker']

    def test_chair_comes_first(self):
        result = minutes_generator._ordered_members(
            self._members, 'Eric Vollheim', 'William Tinker',
        )
        assert result[0] == 'Eric W. Vollheim'

    def test_clerk_comes_last(self):
        result = minutes_generator._ordered_members(
            self._members, 'Eric Vollheim', 'William Tinker',
        )
        assert result[-1] == 'William F. Tinker'

    def test_middle_member_between_chair_and_clerk(self):
        result = minutes_generator._ordered_members(
            self._members, 'Eric Vollheim', 'William Tinker',
        )
        assert result[1] == 'Jeffrey S. Schaaf'

    def test_empty_roles_returns_original_order(self):
        result = minutes_generator._ordered_members(self._members, '', '')
        assert result == self._members

    def test_nickname_matches_full_name(self):
        # 'Jeff Schaaf' (scraped) should match 'Jeffrey S. Schaaf' (members list)
        result = minutes_generator._ordered_members(
            self._members, 'Eric Vollheim', 'Jeff Schaaf',
        )
        assert result[-1] == 'Jeffrey S. Schaaf'


class TestBoardContext:
    def test_includes_member_names(self):
        result = minutes_generator._board_context(BOARD_INFO)
        assert 'Eric W. Vollheim' in result
        assert 'Jeffrey S. Schaaf' in result

    def test_includes_chair_note(self):
        result = minutes_generator._board_context(BOARD_INFO)
        assert 'Eric Vollheim' in result
        assert 'Chair' in result

    def test_empty_board_info(self):
        result = minutes_generator._board_context({})
        assert result == ''


class TestBuildDocx:
    def test_creates_document(self):
        from docx.document import Document as DocxDocument
        doc = minutes_generator._build_docx('2024-03-15', 'Hardwick Town Hall', MOCK_MINUTES_TEXT)
        assert isinstance(doc, DocxDocument)

    def test_document_has_content(self):
        doc = minutes_generator._build_docx('2024-03-15', 'Hardwick Town Hall', MOCK_MINUTES_TEXT)
        full_text = '\n'.join(p.text for p in doc.paragraphs)
        assert 'HARDWICK' in full_text

    def test_handles_empty_content(self):
        doc = minutes_generator._build_docx('2024-03-15', 'Hardwick Town Hall', '')
        assert doc is not None

    def test_allcaps_line_is_bold(self):
        doc = minutes_generator._build_docx('2024-03-15', 'Hardwick Town Hall', 'CALL TO ORDER\nSome prose text.')
        heading_para = next(p for p in doc.paragraphs if p.text == 'CALL TO ORDER')
        assert all(run.bold for run in heading_para.runs if run.text.strip())

    def test_no_blank_paragraphs_in_body_content(self):
        content = 'Line one.\n\n\n\nLine two.'
        doc = minutes_generator._build_docx('2024-03-15', 'Hardwick Town Hall', content)
        # blank separators should only appear in the footer signature block, not body content
        body_blanks = [
            p for p in doc.paragraphs
            if p.text == '' and p != doc.paragraphs[-1]
        ]
        assert len(body_blanks) <= 2  # at most the two footer spacers

    def test_bullet_list_item_stripped(self):
        content = '- Item one\n* Item two'
        doc = minutes_generator._build_docx('2024-03-15', 'Hardwick Town Hall', content)
        texts = [p.text for p in doc.paragraphs]
        assert 'Item one' in texts
        assert 'Item two' in texts
        assert not any(t.startswith('- ') or t.startswith('* ') for t in texts)

    def test_vote_sentence_is_italic(self):
        content = 'Mr. Tinker moved to adjourn. The motion was seconded by Mr. Schaaf and passed unanimously via roll-call vote.'
        doc = minutes_generator._build_docx('2024-03-15', 'Hardwick Town Hall', content)
        vote_para = next(p for p in doc.paragraphs if 'seconded' in p.text)
        assert all(run.italic for run in vote_para.runs if run.text.strip())

    def test_header_includes_location_and_date(self):
        doc = minutes_generator._build_docx('2024-03-15', 'Town Hall', '')
        full_text = '\n'.join(p.text for p in doc.paragraphs)
        assert 'Town Hall' in full_text
        assert '2024-03-15' in full_text

    def test_footer_has_respectfully_submitted(self):
        doc = minutes_generator._build_docx('2024-03-15', 'Hardwick Town Hall', '')
        full_text = '\n'.join(p.text for p in doc.paragraphs)
        assert 'Respectfully submitted' in full_text


class TestGenerateMinutes:
    def test_creates_docx_file(self, tmp_path, mock_openai):
        output = tmp_path / 'minutes_draft_generated.docx'
        result = minutes_generator.generate_minutes(
            MEETING, AGENDA_TEXT, TRANSCRIPT_TEXT, output, BOARD_INFO,
        )
        assert output.exists()
        assert result == output

    def test_calls_llm_twice(self, tmp_path, mock_openai):
        output = tmp_path / 'minutes_draft_generated.docx'
        minutes_generator.generate_minutes(
            MEETING, AGENDA_TEXT, TRANSCRIPT_TEXT, output, BOARD_INFO,
        )
        assert mock_openai.chat.completions.create.call_count == 2

    def test_write_call_includes_date(self, tmp_path, mock_openai):
        output = tmp_path / 'minutes_draft_generated.docx'
        minutes_generator.generate_minutes(
            MEETING, AGENDA_TEXT, TRANSCRIPT_TEXT, output, BOARD_INFO,
        )
        # Second call is _write_minutes; check its user message contains the date
        second_call = mock_openai.chat.completions.create.call_args_list[1]
        messages = second_call.kwargs.get('messages') or second_call[1].get('messages', [])
        user_msg = next(m['content'] for m in messages if m['role'] == 'user')
        assert '2024-03-15' in user_msg

    def test_write_call_includes_agenda_sections(self, tmp_path, mock_openai):
        output = tmp_path / 'minutes_draft_generated.docx'
        minutes_generator.generate_minutes(
            MEETING, AGENDA_TEXT, TRANSCRIPT_TEXT, output, BOARD_INFO,
        )
        second_call = mock_openai.chat.completions.create.call_args_list[1]
        messages = second_call.kwargs.get('messages') or second_call[1].get('messages', [])
        user_msg = next(m['content'] for m in messages if m['role'] == 'user')
        # Agenda sections should appear as ALL-CAPS headings in the skeleton
        assert 'CALL TO ORDER' in user_msg
        assert 'APPROVAL OF MINUTES' in user_msg

    def test_handles_empty_agenda(self, tmp_path, mock_openai):
        output = tmp_path / 'minutes_draft_generated.docx'
        minutes_generator.generate_minutes(
            MEETING, '', TRANSCRIPT_TEXT, output, BOARD_INFO,
        )
        assert output.exists()

    def test_strips_markdown_asterisks(self, tmp_path, monkeypatch):
        mock_client = MagicMock()
        mock_client.chat.completions.create.side_effect = [
            _make_openai_response('facts here'),
            _make_openai_response('**Bold text** should become plain text.'),
        ]
        monkeypatch.setattr(
            'towncommoniq.minutes_generator.OpenAI',
            lambda **kw: mock_client,
        )
        output = tmp_path / 'minutes_draft_generated.docx'
        minutes_generator.generate_minutes(MEETING, '', TRANSCRIPT_TEXT, output)
        doc_text = '\n'.join(p.text for p in __import__('docx').Document(str(output)).paragraphs)
        assert '**' not in doc_text

    def test_strips_content_before_finalized_divider(self, tmp_path, monkeypatch):
        mock_client = MagicMock()
        mock_client.chat.completions.create.side_effect = [
            _make_openai_response('facts here'),
            _make_openai_response(
                'unfilled skeleton text\n\nFinalized version\n\nActual minutes content here.'
            ),
        ]
        monkeypatch.setattr(
            'towncommoniq.minutes_generator.OpenAI',
            lambda **kw: mock_client,
        )
        output = tmp_path / 'minutes.docx'
        minutes_generator.generate_minutes(MEETING, '', TRANSCRIPT_TEXT, output)
        doc_text = '\n'.join(p.text for p in __import__('docx').Document(str(output)).paragraphs)
        assert 'Actual minutes content here.' in doc_text
        assert 'unfilled skeleton' not in doc_text
